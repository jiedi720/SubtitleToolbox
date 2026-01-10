"""
Word文档生成与合并模块
负责将字幕文件转换为Word文档，并提供Word文档合并功能。
"""

import os
import pythoncom
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

try: 
    from docx import Document
    HAS_DOCX = True
except ImportError: 
    HAS_DOCX = False

try: 
    import win32com.client as win32
    HAS_WIN32 = True
except ImportError: 
    HAS_WIN32 = False

# 导入自定义模块
from function.file_utils import get_organized_path, get_save_path, find_files_recursively
from function.parsers import parse_subtitle_to_list
from function.naming import generate_output_name, clean_filename_title
from function.volumes import smart_group_files

def run_word_creation_task(target_dir, log_func, progress_bar, root, batch_size=0, output_dir=None, volume_pattern="智能"):
    """运行Word文档生成任务
    
    从指定目录扫描字幕文件，生成带时间戳的Word文档。
    
    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条信号
        root: 根窗口
        batch_size: 批量大小
        output_dir: 输出目录
        volume_pattern: 分卷模式
    """
    if not HAS_DOCX: 
        return log_func("❌ 错误: 缺少 python-docx 库")
    
    log_func(f"[Word生成] 扫描目录: {target_dir.replace('/', '\\')}", tag="word_blue")
    # 递归查找字幕文件
    files = find_files_recursively(target_dir, ('.srt', '.vtt', '.ass'))
    if not files: 
        return log_func("❌ 未找到字幕。")

    # 智能分组文件
    file_groups = smart_group_files(files, batch_size)
    total_files = len(files)
    count = 0

    # 确定基础输出目录
    base_output_dir = output_dir if output_dir else target_dir

    for group in file_groups:
        if not group: 
            continue
        
        # 生成输出文件名
        out_name = generate_output_name([os.path.basename(f) for f in group], ".docx", volume_pattern)
        # 获取组织化路径
        out_path = get_organized_path(base_output_dir, out_name)
        
        try:
            doc = Document()
            for i, fp in enumerate(group):
                title_text = clean_filename_title(os.path.basename(fp))
                section = doc.sections[0] if i == 0 else doc.add_section()
                section.top_margin = section.bottom_margin = Mm(25)
                section.left_margin = section.right_margin = Mm(25)
                
                # 设置页眉
                header_para = section.header.paragraphs[0]
                header_para.text = title_text
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 添加标题
                doc.add_heading(title_text, level=1)
                
                # 解析字幕内容
                content_list = parse_subtitle_to_list(fp)
                
                if not content_list:
                    doc.add_paragraph("[无对白内容]")
                else:
                    for time_str, text in content_list:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(4)
                        run = p.add_run(f"[{time_str}]  ")
                        run.bold = True
                        p.add_run(text)

                count += 1
                # 更新进度，支持不同类型的进度回调
                try:
                    # 尝试PyQt的信号方式（progress_bar是信号对象）
                    progress_bar.emit(int(count / total_files * 100))
                except AttributeError:
                    try:
                        # 尝试直接调用方式（progress_bar是emit方法本身）
                        progress_bar(int(count / total_files * 100))
                    except Exception as e:
                        pass
            
            # 保存文档
            doc.save(out_path)
            log_func(f"📄 已生成: {os.path.join('word', out_name).replace('/', '\\')}", tag="word_blue")
        except Exception as e: 
            log_func(f"❌ 生成失败: {e}")
    
    # 重置进度条
    try:
        progress_bar.emit(0)
    except AttributeError:
        try:
            progress_bar(0)
        except Exception as e:
            pass

