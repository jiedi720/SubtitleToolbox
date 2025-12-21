import os
import re
import configparser
import tkinter as tk
from tkinter import filedialog, scrolledtext, ttk, colorchooser, font, messagebox, simpledialog
import pysubs2
import pysrt
import threading

# ================= Word 导出依赖 =================
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm
except ImportError:
    pass

# ================= ReportLab (PDF) =================
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Flowable, Frame, PageTemplate, NextPageTemplate
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ================= 全局配置 =================
CONFIG_FILE = "字幕工具箱.ini"
FONT_NAME_BODY = 'Batang'
FONT_PATH_BODY = "C:/Windows/Fonts/batang.ttc"
FONT_NAME_ENG = 'Arial'
FONT_PATH_ENG = "C:/Windows/Fonts/arial.ttf"

try:
    if os.path.exists(FONT_PATH_BODY):
        pdfmetrics.registerFont(TTFont(FONT_NAME_BODY, FONT_PATH_BODY, subfontIndex=0))
    else:
        fallback = "C:/Windows/Fonts/malgun.ttf"
        if os.path.exists(fallback):
            FONT_NAME_BODY = 'MalgunGothic'; pdfmetrics.registerFont(TTFont(FONT_NAME_BODY, fallback))
        else: FONT_NAME_BODY = 'Helvetica'
    if os.path.exists(FONT_PATH_ENG):
        pdfmetrics.registerFont(TTFont(FONT_NAME_ENG, FONT_PATH_ENG))
    else:
        FONT_NAME_ENG = 'Helvetica'
except Exception:
    FONT_NAME_BODY = 'Helvetica'; FONT_NAME_ENG = 'Helvetica'

# ================= PDF 辅助类 =================
class Bookmark(Flowable):
    def __init__(self, key): Flowable.__init__(self); self.key = key
    def draw(self): self.canv.bookmarkPage(self.key)
    def wrap(self, w, h): return (0, 0)

class OutlineEntry(Flowable):
    def __init__(self, title, key): Flowable.__init__(self); self.title = title; self.key = key
    def draw(self): self.canv.addOutlineEntry(self.title, self.key, level=0, closed=True)
    def wrap(self, w, h): return (0, 0)

class TOCFinished(Flowable):
    def wrap(self, w, h): return (0, 0)
    def draw(self): pass

class MyDocTemplate(SimpleDocTemplate):
    def __init__(self, filename, **kw):
        SimpleDocTemplate.__init__(self, filename, **kw)
        self.page_offset = 0  # 初始化偏移量
    
    def afterFlowable(self, flowable):
        if isinstance(flowable, TOCFinished): 
            self.page_offset = self.page  # 目录结束后设置偏移量
        if isinstance(flowable, Paragraph):
            text = flowable.getPlainText()
            if flowable.style.name == 'ChapterTitle':
                key = getattr(flowable, '_bookmarkName', None)
                # 使用相对页码（减去目录页数）
                pg = max(1, self.page - self.page_offset)
                self.notify('TOCEntry', (0, text, pg, key) if key else (0, text, pg))

def footer_content(canvas, doc):
    """绘制页码的页脚函数"""
    physical_pg = canvas.getPageNumber()
    offset = getattr(doc, 'page_offset', 0)
    pg = physical_pg - offset
    
    # 修改后的逻辑：正文页才显示页码（pg > 0）
    if pg > 0:
        canvas.saveState()
        canvas.setFont(FONT_NAME_ENG, 9)
        canvas.drawCentredString(A4[0]/2, 10*mm, str(pg))
        canvas.restoreState()

# ================= 主程序 =================
class UnifiedApp:
    def __init__(self, root):
        self.root = root
        self.root.title("字幕工具箱")
        
        win_w, win_h = 1000, 600
        cx = int(self.root.winfo_screenwidth()/2 - win_w/2)
        cy = int(self.root.winfo_screenheight()/2 - win_h/2)
        self.root.geometry(f'{win_w}x{win_h}+{cx}+{cy}')
        
        self.ui_font = ("Microsoft YaHei", 10)
        self.ui_font_bold = ("Microsoft YaHei", 10, "bold")
        self.ui_font_small = ("Microsoft YaHei", 9)

        self.default_kor_raw = "Style: KOR - Noto Serif KR,Noto Serif KR SemiBold,20,&H0026FCFF,&H000000FF,&H50000000,&H00000000,-1,0,0,0,100,100,0.1,0,1,0.6,0,2,10,10,34,1"
        self.default_chn_raw = "Style: CHN - Drama,小米兰亭,17,&H28FFFFFF,&H000000FF,&H64000000,&H00000000,-1,0,0,0,100,100,0,0,1,0.5,0,2,10,10,15,1"

        self.presets = self.load_presets()
        self.system_fonts = sorted(font.families())
        
        self.path_var = tk.StringVar()
        self.current_preset_name = tk.StringVar(value="默认")
        self.task_mode = tk.StringVar(value="ASS")
        self.do_pdf = tk.BooleanVar(value=True)
        self.do_word = tk.BooleanVar(value=True)

        self.setup_ui()

    def setup_ui(self):
        global_frame = tk.LabelFrame(self.root, text="全局设置", font=self.ui_font, padx=10, pady=5)
        global_frame.pack(fill="x", padx=15, pady=5)

        tk.Button(global_frame, text="📂", command=self.browse_folder, width=3, font=self.ui_font).pack(side="right", padx=5)
        tk.Entry(global_frame, textvariable=self.path_var, font=self.ui_font).pack(side="right", fill="x", expand=True, padx=5)
        tk.Label(global_frame, text="工作目录:", font=self.ui_font_bold).pack(side="right")
        
        tk.Label(global_frame, text="执行任务:", font=self.ui_font_bold).pack(side="left", padx=(0, 10))
        tk.Radiobutton(global_frame, text="合并/转换字幕", variable=self.task_mode, value="ASS", font=self.ui_font).pack(side="left")
        tk.Radiobutton(global_frame, text="生成台词剧本", variable=self.task_mode, value="PDF", font=self.ui_font).pack(side="left", padx=10)

        tk.Label(global_frame, text="剧本格式:", font=self.ui_font_small).pack(side="left", padx=(20, 0))
        tk.Checkbutton(global_frame, text="PDF", variable=self.do_pdf, font=self.ui_font_small).pack(side="left")
        tk.Checkbutton(global_frame, text="Word", variable=self.do_word, font=self.ui_font_small).pack(side="left", padx=5)

        ass_frame = tk.LabelFrame(self.root, text="ASS样式配置 (仅对字幕转换生效)", font=self.ui_font, padx=10, pady=5)
        ass_frame.pack(fill="x", padx=15, pady=5)

        preset_frame = tk.Frame(ass_frame); preset_frame.pack(fill="x", pady=5)
        tk.Label(preset_frame, text="配置方案:", font=self.ui_font_bold).pack(side="left")
        self.preset_combo = ttk.Combobox(preset_frame, textvariable=self.current_preset_name, font=self.ui_font, width=15, state='readonly')
        self.preset_combo['values'] = list(self.presets.keys()); self.preset_combo.pack(side="left", padx=5)
        self.preset_combo.bind("<<ComboboxSelected>>", self.on_preset_change)
        
        tk.Button(preset_frame, text="💾", command=self.save_preset_dialog, bg="#14d6f0", font=self.ui_font).pack(side="left", padx=2, ipadx=4)
        tk.Button(preset_frame, text="❌", command=self.delete_preset, bg="#ef090d", fg="white", font=self.ui_font).pack(side="left", padx=2, ipadx=4)
        tk.Button(preset_frame, text="📝", command=self.open_config_file, bg="#FEE60A", font=self.ui_font).pack(side="left", padx=2, ipadx=4)

        panels_frame = tk.Frame(ass_frame); panels_frame.pack(fill="x")
        curr_data = self.presets.get("默认", {"kor": self.default_kor_raw, "chn": self.default_chn_raw})
        self.kor_parsed = self.parse_ass_style(curr_data["kor"]); self.chn_parsed = self.parse_ass_style(curr_data["chn"])
        self.kor_panel = self.create_style_panel(panels_frame, "外语样式", self.kor_parsed, "안녕하세요 こんにちは Hello")
        self.chn_panel = self.create_style_panel(panels_frame, "中文样式", self.chn_parsed, "你好世界 Hello")

        self.start_btn = tk.Button(self.root, text="开始处理", command=self.start_thread, bg="#0078d7", fg="white", font=self.ui_font_bold, width=20, height=1)
        self.start_btn.pack(pady=10, padx=20)
        self.progress = ttk.Progressbar(self.root, orient="horizontal", mode="determinate"); self.progress.pack(fill="x", padx=20, pady=5)
        self.log_area = scrolledtext.ScrolledText(self.root, height=12, font=("Consolas", 9), state='disabled', bg="#f8f9fa"); self.log_area.pack(fill="both", padx=20, pady=10, expand=False)

    # ================= 辅助方法 =================

    def log(self, message):
        self.log_area.config(state='normal'); self.log_area.insert(tk.END, str(message) + "\n"); self.log_area.see(tk.END); self.log_area.config(state='disabled')

    def browse_folder(self):
        d = filedialog.askdirectory(); 
        if d: self.path_var.set(d)

    def load_presets(self):
        default = {"默认": {"kor": self.default_kor_raw, "chn": self.default_chn_raw}}
        if os.path.exists(CONFIG_FILE):
            try:
                config = configparser.ConfigParser(); config.read(CONFIG_FILE, encoding="utf-8")
                if config.sections(): return {s: {"kor": config.get(s, "kor"), "chn": config.get(s, "chn")} for s in config.sections()}
            except: pass
        return default

    def save_presets_to_file(self):
        try:
            config = configparser.ConfigParser()
            for name, styles in self.presets.items(): config[name] = styles
            with open(CONFIG_FILE, "w", encoding="utf-8") as f: config.write(f)
        except Exception as e: messagebox.showerror("错误", str(e))

    def open_config_file(self):
        if not os.path.exists(CONFIG_FILE): self.save_presets_to_file()
        try: os.startfile(CONFIG_FILE)
        except Exception as e: messagebox.showerror("错误", str(e))

    def on_preset_change(self, event):
        name = self.current_preset_name.get()
        if name in self.presets:
            d = self.presets[name]; self.kor_parsed = self.parse_ass_style(d["kor"]); self.chn_parsed = self.parse_ass_style(d["chn"])
            self.update_panel_ui(self.kor_panel, self.kor_parsed); self.update_panel_ui(self.chn_panel, self.chn_parsed)

    def parse_ass_style(self, style_line):
        if "Style:" not in style_line: return self.parse_ass_style(self.default_kor_raw)
        parts = style_line.replace("Style:", "").strip().split(',')
        while len(parts) < 23: parts.append("0")
        return {"font": parts[1].strip(), "size": parts[2].strip(), "color": parts[3].strip(), "bold": 1 if parts[7].strip() == "-1" else 0, "ml": parts[19].strip(), "mr": parts[20].strip(), "mv": parts[21].strip(), "raw": style_line.strip()}

    def construct_style_line(self, original_raw, ui, style_name):
        parts = original_raw.replace("Style:", "").strip().split(',')
        while len(parts) < 23: parts.append("0")
        parts[0], parts[1], parts[2], parts[3] = style_name, ui["font_var"].get(), ui["size_var"].get(), ui["color_var"].get()
        parts[7] = "-1" if ui["bold_var"].get() else "0"; parts[19], parts[20], parts[21] = ui["ml_var"].get(), ui["mr_var"].get(), ui["mv_var"].get()
        return "Style: " + ",".join(parts)

    def create_style_panel(self, parent, title, initial_data, preview_text):
        panel = tk.LabelFrame(parent, text=title, font=self.ui_font, padx=10, pady=5); panel.pack(side="left", fill="both", expand=True, padx=8)
        vars = {"font_var": tk.StringVar(value=initial_data["font"]), "size_var": tk.StringVar(value=initial_data["size"]), "color_var": tk.StringVar(value=initial_data["color"]), "ml_var": tk.StringVar(value=initial_data["ml"]), "mr_var": tk.StringVar(value=initial_data["mr"]), "mv_var": tk.StringVar(value=initial_data["mv"]), "bold_var": tk.IntVar(value=initial_data["bold"])}
        f_row = tk.Frame(panel); f_row.pack(fill="x", pady=2); tk.Label(f_row, text="字体:", font=self.ui_font, width=4, anchor="w").pack(side="left")
        combo = ttk.Combobox(f_row, textvariable=vars["font_var"], values=self.system_fonts, state="readonly", font=self.ui_font); combo.pack(side="left", fill="x", expand=True)
        p_label = tk.Label(panel, text=preview_text, font=("Arial", 10), bg="white", relief="sunken", height=1); p_label.pack(fill="x", pady=(2, 8))
        def update_p(*args):
            try: f, sz = vars["font_var"].get(), int(float(vars["size_var"].get())); w = "bold" if vars["bold_var"].get() else "normal"; p_label.config(font=(f, min(sz, 28), w))
            except: pass
        combo.bind("<<ComboboxSelected>>", update_p)
        param_row = tk.Frame(panel); param_row.pack(fill="x", pady=2); left = tk.Frame(param_row); left.pack(side="left")
        tk.Label(left, text="号:", font=self.ui_font).pack(side="left")
        tk.Spinbox(left, from_=1, to=200, textvariable=vars["size_var"], width=3, font=self.ui_font, command=update_p).pack(side="left", padx=(0,5))
        tk.Checkbutton(left, text="B", variable=vars["bold_var"], font=("Georgia", 10, "bold"), command=update_p).pack(side="left", padx=(0,5))
        c_btn = tk.Label(left, width=2, relief="ridge", borderwidth=1); c_btn.pack(side="left", padx=(2,2), fill="y")
        c_btn.bind("<Button-1>", lambda e: self.pick_color(vars["color_var"], c_btn))
        tk.Entry(left, textvariable=vars["color_var"], font=self.ui_font_small, width=11).pack(side="left")
        self.sync_color(vars["color_var"], c_btn); tk.Label(param_row, text="").pack(side="left", expand=True); right = tk.Frame(param_row); right.pack(side="right")
        def qs(txt, v): tk.Label(right, text=txt, font=self.ui_font_small).pack(side="left", padx=(3,0)); tk.Spinbox(right, from_=0, to=500, textvariable=v, width=3, font=self.ui_font).pack(side="left")
        qs("L:", vars["ml_var"]); qs("R:", vars["mr_var"]); qs("V:", vars["mv_var"])
        update_p(); vars.update({"c_btn": c_btn, "update_func": update_p}); return vars

    def update_panel_ui(self, ui, data):
        ui["font_var"].set(data["font"]); ui["size_var"].set(data["size"]); ui["color_var"].set(data["color"])
        ui["ml_var"].set(data["ml"]); ui["mr_var"].set(data["mr"]); ui["mv_var"].set(data["mv"]); ui["bold_var"].set(data["bold"])
        self.sync_color(ui["color_var"], ui["c_btn"]); ui["update_func"]()

    def sync_color(self, var, btn):
        m = re.search(r'&H[0-9A-F]{2}([0-9A-F]{2})([0-9A-F]{2})([0-9A-F]{2})', var.get().upper())
        if m: btn.config(bg=f"#{m.groups()[2]}{m.groups()[1]}{m.groups()[0]}")

    def pick_color(self, var, btn):
        code = colorchooser.askcolor()
        if code[1]:
            rgb = code[0]; m = re.search(r'&H([0-9A-F]{2})', var.get().upper()); alpha = m.group(1) if m else "00"
            var.set(f"&H{alpha}{int(rgb[2]):02X}{int(rgb[1]):02X}{int(rgb[0]):02X}&"); btn.config(bg=code[1])

    def start_thread(self): threading.Thread(target=self.process, daemon=True).start()

    # ================= 业务逻辑 =================

    def clean_subtitle_text_ass(self, t): 
        return re.sub(r'\[.*?\]|\(.*?\)|{.*?}', '', t).replace('\n', ' ').replace('\\N', ' ').strip()
    
    def clean_subtitle_text_pdf(self, t):
        # 1. 移除标签和换行
        text = re.sub(r'\[.*?\]|\(.*?\)|{.*?}', '', t).replace('\\N', ' ').replace('\n', ' ')
        # 2. 移除不可见 Unicode 控制符
        text = "".join(ch for ch in text if ch.isprintable())
        # 3. 合并空白
        text = re.sub(r'\s+', ' ', text).strip()
        
        # --- 新增：过滤无效行 ---
        # 如果清洗后只剩下标点符号、连字符或空格（例如 "- -", ". .", "---"），则视为空行
        if not text or re.fullmatch(r'[\s\-\.\,\?\!\_]+', text):
            return ""
            
        return text

    # --- 剧名保留逻辑：去掉年份及之后的所有内容 ---
    def clean_filename_title(self, n):
        # 1. 移除扩展名和音轨/字幕提取后缀 (处理 _track19 或 _6_text)
        name = os.path.splitext(n)[0]
        name = re.sub(r'(_track\d+|_?\d+_text).*$', '', name, flags=re.IGNORECASE)
        
        # 2. 关键修复：先移除剧名中间夹杂的年份 (19xx 或 20xx)
        # 这样即便年份在 S01E01 前面，也不会触发后续的“切断”逻辑
        name = re.sub(r'[. ](19|20)\d{2}([. ]|$)', '.', name)

        # 3. 定义切断关键词（这些词之后的内容全部舍弃）
        # 移除了 \d{4}，改为在步骤2中精准抽离，防止误切
        stop_keywords = [
            r'Episode', r'2160p', r'1080p', r'720p', r'MULTi', 
            r'NF', r'Netflix', r'WEB-DL', r'WEBrip', r'BluRay', 
            r'DDP5', r'Atmos', r'x265', r'x264', r'CMCTV', r'ARiC'
        ]
        
        # 构建正则：匹配这些词汇及其后面的一切
        pattern = r'\.(' + '|'.join(stop_keywords) + r')\..*$'
        name = re.sub(pattern, '', name, flags=re.IGNORECASE)
        
        # 兜底清理：处理关键词直接连着点号结尾的情况
        for kw in [r'Episode', r'MULTi', r'NF', r'CMCTV']:
            name = re.sub(r'\.' + kw + r'.*$', '', name, flags=re.IGNORECASE)

        # 4. 最终美化：点号转空格，处理连续多余空格
        clean_name = name.replace('.', ' ')
        return re.sub(r'\s+', ' ', clean_name).strip()

    def generate_pdf_name(self, files):
        first = files[0]; m = re.search(r'^(.*?)[. ]S\d+E\d+', first, re.IGNORECASE)
        if m:
            series = m.group(1); ss = sorted(list(set([int(re.search(r'[. ]S(\d+)E\d+', f, re.IGNORECASE).group(1)) for f in files if re.search(r'[. ]S(\d+)E\d+', f, re.IGNORECASE)])))
            return f"{series}.S{ss[0]:02d}.pdf" if len(ss)==1 else f"{series}.S{ss[0]:02d}-{ss[-1]:02d}.pdf"
        return os.path.splitext(first)[0] + ".pdf"

    def process(self):
        target_dir = self.path_var.get().strip()
        if not target_dir or not os.path.exists(target_dir): messagebox.showerror("错误", "无效的目录"); return
        self.start_btn.config(state='disabled'); self.log(f"--- 任务启动: {target_dir} ---")
        try:
            mode = self.task_mode.get()
            if mode == "ASS":
                self.process_ass(target_dir)
            elif mode == "PDF":
                if not self.do_pdf.get() and not self.do_word.get():
                    self.log("⚠️ 未勾选任何导出格式。")
                else:
                    if self.do_pdf.get(): self.process_pdf(target_dir)
                    if self.do_word.get(): self.process_word(target_dir)
            self.log("✅ 任务完成！")
        except Exception as e: self.log(f"❌ 严重错误: {e}")
        finally: self.start_btn.config(state='normal'); self.progress["value"] = 0

    def process_ass(self, d):
        l_k = self.construct_style_line(self.kor_parsed["raw"], self.kor_panel, "KOR - Noto Serif KR"); l_c = self.construct_style_line(self.chn_parsed["raw"], self.chn_panel, "CHN - Drama")
        hdr = f"[Script Info]\nScriptType: v4.00+\nWrapStyle: 0\nScaledBorderAndShadow: yes\n\n[V4+ Styles]\nFormat: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding\n{l_k}\n{l_c}\n\n[Events]\nFormat: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text"
        all_f = os.listdir(d); duals = [f for f in all_f if f.lower().endswith('.dual.srt')]; srts = [f for f in all_f if f.lower().endswith('.srt') and '.dual.' not in f.lower()]; tasks = []; done = set()
        for f in duals: tasks.append({"type": "convert", "path": os.path.join(d, f), "name": f}); m = re.search(r'[Ss]\d{2}[Ee]\d{2}', f); 
        if m: done.add(m.group().upper())
        gps = {}
        for f in srts:
            m = re.search(r'[Ss]\d{2}[Ee]\d{2}', f); 
            if m:
                ep = m.group().upper()
                if ep not in done: gps.setdefault(ep, []).append(f)
        for ep, fl in gps.items():
            chi = [f for f in fl if any(x in f.lower() for x in ['chi', 'chs', 'cht'])]; oth = [f for f in fl if f not in chi]
            if chi and oth: tasks.append({"type": "merge", "ep": ep, "chi": os.path.join(d, chi[0]), "other": os.path.join(d, oth[0]), "base": oth[0]})
        self.progress["maximum"] = len(tasks)
        for i, t in enumerate(tasks):
            evs = []
            if t["type"] == "convert":
                subs = pysubs2.load(t["path"]); split = 0
                for j in range(1, len(subs)):
                    if subs[j].start < subs[j-1].start: split = j; break
                for j, l in enumerate(subs):
                    sty = "KOR - Noto Serif KR" if j < split else "CHN - Drama"
                    st, et = pysubs2.time.ms_to_str(l.start, fractions=True).replace(',','.')[:-1], pysubs2.time.ms_to_str(l.end, fractions=True).replace(',','.')[:-1]
                    evs.append(f"Dialogue: 0,{st},{et},{sty},,0,0,0,,{l.text.replace('\\N', ' ').strip()}")
                out = t["name"].replace('.dual.srt', '.ass')
            else:
                s1, s2 = pysubs2.load(t["other"]), pysubs2.load(t["chi"])
                for l in s1:
                    c = self.clean_subtitle_text_ass(l.text)
                    if c:
                        st, et = pysubs2.time.ms_to_str(l.start, fractions=True).replace(',','.')[:-1], pysubs2.time.ms_to_str(l.end, fractions=True).replace(',','.')[:-1]
                        evs.append(f"Dialogue: 0,{st},{et},KOR - Noto Serif KR,,0,0,0,,{c}")
                for l in s2:
                    c = self.clean_subtitle_text_ass(l.text)
                    if c:
                        st, et = pysubs2.time.ms_to_str(l.start, fractions=True).replace(',','.')[:-1], pysubs2.time.ms_to_str(l.end, fractions=True).replace(',','.')[:-1]
                        evs.append(f"Dialogue: 0,{st},{et},CHN - Drama,,0,0,0,,{c}")
                out = re.split(r'_track\d+', t["base"], flags=re.IGNORECASE)[0].rstrip('._ ') + ".ass"
            with open(os.path.join(d, out), 'w', encoding='utf-8-sig') as f: f.write(hdr + "\n".join(evs))
            self.progress["value"] = i+1; self.root.update_idletasks()

    def process_pdf(self, target_dir):
        files = [f for f in os.listdir(target_dir) if f.lower().endswith('.srt')]; files.sort()
        if not files: self.log("[PDF] 未找到 SRT 文件。"); return
        out_name = self.generate_pdf_name(files); out_path = os.path.join(target_dir, out_name)
        
        # 1. 初始化文档
        doc = MyDocTemplate(out_path, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')
        
        # 2. 严格定义模板回调
        doc.addPageTemplates([
            PageTemplate(id='TOC', frames=frame, onPage=lambda c,d: None),
            PageTemplate(id='Body', frames=frame, onPage=footer_content)
        ])
        
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle('ChapterTitle', parent=styles['Heading1'], fontName=FONT_NAME_BODY, fontSize=16, leading=20, spaceAfter=10, textColor=colors.darkblue)
        toc_h = ParagraphStyle('TOCHeader', parent=styles['Heading1'], fontName=FONT_NAME_ENG, fontSize=20, leading=22, spaceAfter=20, alignment=TA_CENTER)
        body = ParagraphStyle('SubtitleBody', parent=styles['Normal'], fontName=FONT_NAME_BODY, fontSize=10, leading=14, spaceAfter=4)
        
        # 3. 构建目录部分
        story = [
            Bookmark("TOC"), 
            OutlineEntry("Content", "TOC"), 
            Paragraph("<b>Content</b>", toc_h)
        ]
        
        toc = TableOfContents(); toc.dotsMinLevel = 0
        toc.levelStyles = [ParagraphStyle('TOC1', fontName=FONT_NAME_BODY, fontSize=12, leftIndent=20, firstLineIndent=-20, spaceBefore=5)]
        
        story.append(toc)
        story.append(TOCFinished()) # 计算偏移量标记
        
        # --- 修复关键点：先切换模板，再执行分页 ---
        story.append(NextPageTemplate('Body')) 
        story.append(PageBreak()) 
        
        # 4. 构建正文内容
        self.progress["maximum"] = len(files)
        for i, fname in enumerate(files):
            dname = self.clean_filename_title(fname)
            key = f"CH_{i}"
            
            # 添加书签和目录条目
            story.append(Bookmark(key))
            story.append(OutlineEntry(dname, key))
            
            # 章节标题
            p = Paragraph(dname, h1)
            p._bookmarkName = key
            story.append(p)
            story.append(Spacer(1, 10))
            
            try: subs = pysrt.open(os.path.join(target_dir, fname), encoding='utf-8')
            except: subs = pysrt.open(os.path.join(target_dir, fname), encoding='gbk')
            
            for s in subs:
                txt = self.clean_subtitle_text_pdf(s.text)
                if not txt: continue
                safe = txt.replace('<','&lt;').replace('>','&gt;')
                story.append(Paragraph(f"<b>[{s.start}]</b>  {safe}", body))
            
            # 章节间分页
            if i < len(files)-1:
                story.append(PageBreak())
            
            self.progress["value"] = i + 1
            self.root.update_idletasks()
            
        # 5. 执行两次编译以生成目录和页码
        doc.multiBuild(story)
        self.log(f"[PDF] 生成完毕: {out_name}")

    def process_word(self, target_dir):
        try:
            from docx import Document
        except ImportError:
            messagebox.showerror("依赖缺失", "请运行 pip install python-docx 以使用此功能")
            return

        files = [f for f in os.listdir(target_dir) if f.lower().endswith('.srt')]; files.sort()
        if not files: self.log("[Word] 未找到 SRT 文件。"); return
        out_name = self.generate_pdf_name(files).replace(".pdf", ".docx")
        out_path = os.path.join(target_dir, out_name)
        doc = Document()
        for section in doc.sections:
            section.top_margin = Mm(25); section.bottom_margin = Mm(25); section.left_margin = Mm(25); section.right_margin = Mm(25)

        self.progress["maximum"] = len(files)
        for i, fname in enumerate(files):
            dname = self.clean_filename_title(fname)
            self.log(f"[Word] 正在写入: {dname}")
            heading = doc.add_heading(dname, level=1)
            run = heading.runs[0]; run.font.color.rgb = RGBColor(0, 51, 102)

            try: subs = pysrt.open(os.path.join(target_dir, fname), encoding='utf-8')
            except: subs = pysrt.open(os.path.join(target_dir, fname), encoding='gbk')

            for s in subs:
                txt = self.clean_subtitle_text_pdf(s.text)
                if not txt or len(txt.strip()) == 0: continue # 严格跳过空行
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
                time_run = p.add_run(f"[{s.start}]  "); time_run.bold = True; time_run.font.size = Pt(10)
                text_run = p.add_run(txt); text_run.font.size = Pt(10)

            if i < len(files)-1: doc.add_page_break()
            self.progress["value"] = i + 1; self.root.update_idletasks()
        doc.save(out_path); self.log(f"[Word] 生成完毕: {out_name}")

    def save_preset_dialog(self):
        n = simpledialog.askstring("保存方案", "请输入名称:", initialvalue=self.current_preset_name.get())
        if n:
            l_k = self.construct_style_line(self.kor_parsed["raw"], self.kor_panel, "KOR"); l_c = self.construct_style_line(self.chn_parsed["raw"], self.chn_panel, "CHN")
            self.presets[n.strip()] = {"kor": l_k, "chn": l_c}; self.save_presets_to_file()
            self.preset_combo['values'] = list(self.presets.keys()); self.preset_combo.set(n.strip())

    def delete_preset(self):
        n = self.current_preset_name.get()
        if n != "默认" and messagebox.askyesno("删除", f"确定删除 [{n}] 吗？"):
            del self.presets[n]; self.save_presets_to_file(); self.current_preset_name.set("默认"); self.preset_combo['values'] = list(self.presets.keys()); self.on_preset_change(None)

if __name__ == "__main__":
    root = tk.Tk(); app = UnifiedApp(root); root.mainloop()