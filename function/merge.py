"""
文件合并模块
负责各种文档类型的合并功能，包括PDF、TXT和Word文档。
"""

import os
import re

__all__ = [
    'run_pdf_merge_task',
    'run_txt_merge_task',
    'run_md_merge_task',
    'run_docx_merge_task',
    'run_win32_merge_task',
    'execute_merge_tasks'
]

# PDF合并所需的导入
try:
    from pypdf import PdfWriter as PdfMerger 
except ImportError:
    PdfMerger = None

# Word合并所需的导入
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 尝试导入pywin32，用于更高效的Word合并
try:
    import pythoncom
    import win32com.client as win32
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False


def run_pdf_merge_task(target_dir, log_func, progress_bar, root, output_dir=None, stop_flag=False):
    """运行PDF文档合并任务
    
    合并多个PDF文档为一个。
    
    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条信号
        root: 根窗口
        output_dir: 输出目录
    """
    if PdfMerger is None: 
        return log_func("❌ 缺少 pypdf 库，请安装。")
    
    # 查找PDF文件
    root_files = sorted([os.path.join(target_dir, f) for f in os.listdir(target_dir) 
                        if f.lower().endswith('.pdf') and "合并" not in f])
    
    target_files = root_files if root_files else []
    save_dir = target_dir

    if not target_files:
        # 检查目标目录下的pdf子文件夹
        sub_dir = os.path.join(target_dir, "pdf")
        if os.path.exists(sub_dir):
            target_files = sorted([os.path.join(sub_dir, f) for f in os.listdir(sub_dir)
                                 if f.lower().endswith('.pdf') and "合并" not in f])
            save_dir = sub_dir

    if not target_files: 
        return log_func("❌ 未找到 PDF 文件")

    merger = PdfMerger()
    try:
        for i, f in enumerate(target_files):
                # 检查停止标志
                if stop_flag[0]:
                    break
                log_func(f"合并中: {os.path.basename(f)}")
                merger.append(f)
                # 更新进度，支持不同类型的进度回调
                try:
                    # 尝试PyQt的信号方式（progress_bar是信号对象）
                    progress_bar.emit(int((i + 1) / len(target_files) * 100))
                except AttributeError:
                    try:
                        # 尝试直接调用方式（progress_bar是emit方法本身）
                        progress_bar(int((i + 1) / len(target_files) * 100))
                    except Exception as e:
                        pass
            
        # 检查是否停止，停止则不输出文件
        if stop_flag[0]:
            merger.close()
            log_func("⚠️ 任务已停止，未生成合并文件")
        else:
            # 输出合并后的PDF
            out_path = os.path.join(save_dir, "PDF合并.pdf")
            merger.write(out_path)
            merger.close()
            log_func(f"✅ 合并成功: {out_path.replace('/', '\\')}")
    except Exception as e: 
        log_func(f"❌ 错误: {e}")
    finally: 
        # 重置进度条
        try:
            progress_bar.emit(0)
        except AttributeError:
            try:
                progress_bar(0)
            except Exception as e:
                pass


def run_txt_merge_task(target_dir, log_func, progress_bar, root, output_dir=None, stop_flag=False):
    """运行TXT文档合并任务

    合并多个TXT文档为一个。

    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条信号
        root: 根窗口
        output_dir: 输出目录
    """
    # 查找TXT文件
    root_files = sorted([os.path.join(target_dir, f) for f in os.listdir(target_dir) 
                        if f.lower().endswith('.txt') and "合并" not in f])
    
    target_files = []
    save_dir = target_dir

    if root_files:
        target_files = root_files
    else:
        # 检查目标目录下的txt子文件夹
        sub_dir = os.path.join(target_dir, "txt")
        if os.path.exists(sub_dir):
            sub_files = sorted([os.path.join(sub_dir, f) for f in os.listdir(sub_dir)
                               if f.lower().endswith('.txt') and "合并" not in f])
            if sub_files:
                target_files = sub_files
                save_dir = sub_dir

    if not target_files:
        return log_func("❌ 未找到 TXT 文件")

    total = len(target_files)
    out_path = os.path.join(save_dir, "TXT合并.txt")

    try:
        with open(out_path, 'w', encoding='utf-8') as outfile:
            for i, fp in enumerate(target_files):
                # 检查停止标志
                if stop_flag[0]:
                    break
                log_func(f"合并中: {os.path.basename(fp)}")
                with open(fp, 'r', encoding='utf-8') as infile:
                    outfile.write(infile.read())
                    outfile.write("\n\n" + "="*50 + "\n\n") 
                # 更新进度，支持不同类型的进度回调
                try:
                    # 尝试PyQt的信号方式（progress_bar是信号对象）
                    progress_bar.emit(int((i + 1) / total * 100))
                except AttributeError:
                    try:
                        # 尝试直接调用方式（progress_bar是emit方法本身）
                        progress_bar(int((i + 1) / total * 100))
                    except Exception as e:
                        pass
            
        # 检查是否停止
        if stop_flag[0]:
            log_func("⚠️ 任务已停止，未生成完整合并文件")
            # 删除不完整的文件
            if os.path.exists(out_path):
                os.remove(out_path)
        else:
            log_func(f"✅ 合并成功: {out_path.replace('/', '\\')}")
    except Exception as e:
        log_func(f"❌ 合并失败: {e}")
    finally:
        # 重置进度条
        try:
            progress_bar.emit(0)
        except AttributeError:
            try:
                progress_bar(0)
            except Exception as e:
                pass


def run_md_merge_task(target_dir, log_func, progress_bar, root, output_dir=None, stop_flag=False):
    """运行Markdown文档合并任务

    合并多个Markdown文档为一个。

    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条信号
        root: 根窗口
        output_dir: 输出目录
    """
    # 查找Markdown文件
    root_files = sorted([os.path.join(target_dir, f) for f in os.listdir(target_dir) 
                        if f.lower().endswith('.md') and "合并" not in f])
    
    target_files = []
    save_dir = target_dir

    if root_files:
        target_files = root_files
    else:
        # 检查目标目录下的md子文件夹
        sub_dir = os.path.join(target_dir, "md")
        if os.path.exists(sub_dir):
            sub_files = sorted([os.path.join(sub_dir, f) for f in os.listdir(sub_dir)
                               if f.lower().endswith('.md') and "合并" not in f])
            if sub_files:
                target_files = sub_files
                save_dir = sub_dir

    if not target_files:
        return log_func("❌ 未找到 Markdown 文件")

    total = len(target_files)
    out_path = os.path.join(save_dir, "Markdown合并.md")

    try:
        with open(out_path, 'w', encoding='utf-8') as outfile:
            for i, fp in enumerate(target_files):
                # 检查停止标志
                if stop_flag[0]:
                    break
                log_func(f"合并中: {os.path.basename(fp)}")
                with open(fp, 'r', encoding='utf-8') as infile:
                    outfile.write(infile.read())
                    outfile.write("\n\n---\n\n") 
                # 更新进度，支持不同类型的进度回调
                try:
                    # 尝试PyQt的信号方式（progress_bar是信号对象）
                    progress_bar.emit(int((i + 1) / total * 100))
                except AttributeError:
                    try:
                        # 尝试直接调用方式（progress_bar是emit方法本身）
                        progress_bar(int((i + 1) / total * 100))
                    except Exception as e:
                        pass
            
        # 检查是否停止
        if stop_flag[0]:
            log_func("⚠️ 任务已停止，未生成完整合并文件")
            # 删除不完整的文件
            if os.path.exists(out_path):
                os.remove(out_path)
        else:
            log_func(f"✅ 合并成功: {out_path.replace('/', '\\')}")
    except Exception as e:
        log_func(f"❌ 合并失败: {e}")
    finally:
        # 重置进度条
        try:
            progress_bar.emit(0)
        except AttributeError:
            try:
                progress_bar(0)
            except Exception as e:
                pass


def run_docx_merge_task(target_dir, log_func, progress_bar, root, output_dir=None, stop_flag=False):
    """使用python-docx合并Word文档
    
    通过纯Python方式合并多个Word文档为一个，不依赖Windows COM接口。
    优先保证合并速度和稳定性，而非完美的格式保留。
    
    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条对象
        root: 根窗口
        output_dir: 输出目录
    """
    if not HAS_DOCX:
        return False, "未安装 python-docx 库"
    
    # 查找Word文件 - 极简版本
    log_func(f"🔍 快速扫描目录: {target_dir}")
    
    # 极速文件查找：只扫描根目录，优先保证速度
    target_files = []
    for f in os.listdir(target_dir):
        if f.lower().endswith('.docx') and "~$" not in f and "合并" not in f:
            target_files.append(os.path.join(target_dir, f))
    
    total_files = len(target_files)
    
    if not target_files:
        log_func(f"❌ 未找到 Word 文件")
        return False, "未找到 Word 文件"
    
    log_func(f"📋 找到 {total_files} 个Word文件，开始快速合并...")
    
    # 直接使用当前目录作为保存目录
    out_path = os.path.join(target_dir, "Word合并.docx")
    
    try:
        # 创建新文档
        merged_doc = Document()
        
        # 处理每个文件 - 极简版本
        for i, fp in enumerate(target_files):
            # 检查停止标志
            if stop_flag[0]:
                break
                
            filename = os.path.basename(fp)
            log_func(f"🔄 处理: {filename}")
            
            # 立即更新进度，确保用户看到反馈
            try:
                progress = int((i / total_files) * 100)
                progress_bar.emit(progress)
            except Exception:
                try:
                    progress_bar.set(i / total_files)
                    if root:
                        root.update_idletasks()
                except Exception:
                    try:
                        progress_bar(progress)
                    except Exception:
                        pass
            
            # 极速文件处理：跳过错误检查，直接尝试处理
            try:
                # 打开源文档
                src_doc = Document(fp)
                
                # 极速段落处理：只复制文本，不处理样式
                for para in src_doc.paragraphs:
                    if para.text.strip():
                        merged_doc.add_paragraph(para.text)
                
                # 跳过表格处理，大幅提高速度
                if len(src_doc.tables) > 0:
                    log_func(f"⚠️  跳过 {len(src_doc.tables)} 个表格，如需表格支持请使用其他工具")
                
                # 插入分页符
                if i < total_files - 1:
                    merged_doc.add_page_break()
                    
                log_func(f"✅ 完成: {filename}")
            
            except Exception as e:
                log_func(f"❌ 处理失败: {filename} - {e}")
                continue
        
        # 检查是否停止
        if stop_flag[0]:
            log_func("⚠️ 任务已停止，未生成合并文件")
            return False, "任务已停止"
        
        # 最终进度更新
        try:
            progress_bar.emit(100)
        except Exception:
            try:
                progress_bar.set(1.0)
                if root:
                    root.update_idletasks()
            except Exception:
                try:
                    progress_bar(100)
                except Exception:
                    pass
        
        # 快速保存
        log_func(f"💾 保存: {out_path}")
        merged_doc.save(out_path)
        log_func(f"✅ 合并完成: {out_path}")
        return True, ""
    
    except PermissionError:
        log_func(f"❌ 文件被占用，请关闭 Word 后重试")
        return False, "文件被占用"
    except Exception as e:
        log_func(f"❌ 合并失败: {e}")
        return False, str(e)


def run_win32_merge_task(target_dir, log_func, progress_bar, root, output_dir=None, stop_flag=False):
    """使用pywin32合并Word文档
    
    通过Windows COM接口合并多个Word文档为一个，保留完整格式和表格。
    
    Args:
        target_dir: 目标目录
        log_func: 日志记录函数
        progress_bar: 进度条对象
        root: 根窗口
        output_dir: 输出目录
    """
    if not HAS_WIN32:
        log_func("❌ 错误: 缺少 pywin32 库")
        return
    
    # 初始化COM
    pythoncom.CoInitialize()
    
    # 查找Word文件
    log_func(f"🔍 扫描目录: {target_dir}")
    
    root_files = []
    for f in os.listdir(target_dir):
        if f.lower().endswith('.docx') and "~$" not in f and "合并" not in f:
            root_files.append(os.path.join(target_dir, f))
    
    target_files = root_files if root_files else []
    save_dir = target_dir

    if not target_files:
        # 适配新的分类层级：检测 script/word
        sub_dir = os.path.join(target_dir, "script", "word")
        if os.path.exists(sub_dir):
            target_files = []
            for f in os.listdir(sub_dir):
                if f.lower().endswith('.docx') and "~$" not in f and "合并" not in f:
                    target_files.append(os.path.join(sub_dir, f))
            save_dir = sub_dir

    if not target_files:
        pythoncom.CoUninitialize()
        return log_func("❌ 未找到 Word 文件")

    word = None
    try:
        log_func(f"📋 找到 {len(target_files)} 个Word文件，准备合并...")
        
        # 启动Word应用
        word = win32.Dispatch('Word.Application')
        word.Visible = False
        
        # 创建新文档
        new_doc = word.Documents.Add()
        sel = word.Selection
        
        # 合并文件
        total_files = len(target_files)
        for i, fp in enumerate(target_files):
            # 检查停止标志
            if stop_flag[0]:
                break
                
            filename = os.path.basename(fp)
            # 移除文件扩展名
            filename_no_ext = os.path.splitext(filename)[0]
            log_func(f"🔄 合并中: {filename} ({i+1}/{total_files})")
            
            # 处理文件名，去掉前面的序号（如 "01.【第 Ⅰ 語基】" → "【第 Ⅰ 語基】"）
            # 匹配开头的数字、点和可能的空格，提取后面的内容
            match = re.match(r'^\d+\.\s*(.*)', filename_no_ext)
            if match:
                clean_filename = match.group(1)
            else:
                clean_filename = filename_no_ext
            
            # 插入清理后的文件名作为Heading 1
            sel.TypeText(clean_filename)
            sel.Style = "Heading 1"  # 应用Heading 1样式
            sel.TypeParagraph()  # 插入段落标记
            
            # 插入文件内容
            sel.InsertFile(os.path.abspath(fp))
            
            # 插入分页符
            if i < total_files - 1:
                sel.InsertBreak(Type=7)  # wdPageBreak = 7
            
            # 更新进度
            try:
                progress = int((i + 1) / total_files * 100)
                progress_bar.emit(progress)
            except AttributeError:
                try:
                    progress_bar.set((i + 1) / total_files)
                    if root:
                        root.update_idletasks()
                except AttributeError:
                    try:
                        progress_bar(progress)
                    except AttributeError:
                        pass
        
        # 检查是否停止，停止则不保存文件
        if stop_flag[0]:
            new_doc.Close(SaveChanges=0)  # 0 = wdDoNotSaveChanges
            log_func("⚠️ 任务已停止，未生成合并文件")
        else:
            # 保存合并后的文档
            out_path = os.path.join(save_dir, "Word合并.docx")
            new_doc.SaveAs2(os.path.abspath(out_path), FileFormat=12)  # 12 = wdFormatXMLDocument
            new_doc.Close()
            log_func(f"✅ 合并完成: {out_path}")
    except Exception as e:
        log_func(f"❌ 运行错误: {e}")
    finally:
        # 清理资源
        if word:
            word.Quit()
        pythoncom.CoUninitialize()
        # 重置进度条
        try:
            progress_bar.emit(0)
        except AttributeError:
            try:
                progress_bar.set(0)
                if root:
                    root.update_idletasks()
            except AttributeError:
                try:
                    progress_bar(0)
                except AttributeError:
                    pass


def execute_merge_tasks(path_var, output_path_var, log_callback, update_progress, root, gui, stop_flag=False):
    """
    执行合并任务
    
    Args:
        path_var: 源目录路径
        output_path_var: 输出目录路径
        log_callback: 日志回调函数
        update_progress: 进度回调函数
        root: 根窗口对象
        gui: GUI 对象
    """
    from PySide6.QtWidgets import QMessageBox
    
    # 获取目标目录
    target_dir = path_var.strip()
    if not target_dir or not os.path.exists(target_dir):
        QMessageBox.critical(root, "错误", "请选择有效目录")
        return
    
    # 获取输出目录
    if output_path_var.strip():
        final_out = output_path_var.strip()
    else:
        final_out = target_dir
    
    # 检查Merge标签页中的复选框状态
    try:
        # 获取Merge标签页中的复选框状态
        merge_pdf = gui.MergePDF.isChecked()
        merge_word = gui.MergeWord.isChecked()
        merge_txt = gui.MergeTxt.isChecked()
        
        # 检查是否至少选中了一个合并选项
        if not merge_pdf and not merge_word and not merge_txt:
            log_callback("❌ 请至少选择一个合并选项")
            return
        
        # 根据选中的选项执行相应的合并任务
        if merge_pdf:
            run_pdf_merge_task(
                target_dir, 
                log_callback, 
                update_progress, 
                root, 
                output_dir=final_out,
                stop_flag=stop_flag
            )
        
        if merge_word:
            run_win32_merge_task(
                target_dir, 
                log_callback, 
                update_progress, 
                root, 
                output_dir=final_out,
                stop_flag=stop_flag
            )
        
        if merge_txt:
            run_txt_merge_task(
                target_dir, 
                log_callback, 
                update_progress, 
                root, 
                output_dir=final_out,
                stop_flag=stop_flag
            )
            
    except Exception as e:
        log_callback(f"❌ Merge模式处理失败: {e}")
